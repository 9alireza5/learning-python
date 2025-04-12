print('this code, convers persian numbers in english and sum them')
from docx import Document
file_path = 'you should enter file path here'
doc = Document(file_path)
translation_table = str.maketrans("۰۱۲۳۴۵۶۷۸۹٫", "0123456789.")
for paragraph in doc.paragraphs:
    converted_text = paragraph.text.translate(translation_table)
    print(converted_text)
total_sum = 0
current_number = ''
for para in doc.paragraphs:
    converted_text = para.text.translate(translation_table)
    for char in converted_text:
        if char in '0123456789.':
            current_number += char
        else:
            if current_number and current_number != '.':
                total_sum += float(current_number)
            current_number = ''
if current_number and current_number != '.':
    total_sum += float(current_number)
print('total sum of digits in docx file is equal to:',total_sum)